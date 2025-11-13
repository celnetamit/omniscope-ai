"""
Report Generators for PDF, DOCX, and LaTeX formats
"""

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pylatex import Document as LaTeXDocument, Section, Subsection, Command, Package
from pylatex.utils import NoEscape, italic, bold
from typing import Dict, Any, List
from pathlib import Path
from datetime import datetime
import os
from html.parser import HTMLParser
import re

class PDFGenerator:
    """
    PDF report generator using ReportLab
    
    Requirements: 5.1
    """
    
    def __init__(self):
        """Initialize PDF generator"""
        self.styles = getSampleStyleSheet()
        self._create_custom_styles()
    
    def _create_custom_styles(self):
        """Create custom paragraph styles"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        
        # Section heading style
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        ))
        
        # Body text style
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            fontName='Helvetica'
        ))
        
        # Metadata style
        self.styles.add(ParagraphStyle(
            name='Metadata',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#7f8c8d'),
            spaceAfter=20,
            alignment=TA_CENTER
        ))
    
    def _add_header_footer(self, canvas, doc):
        """Add header and footer to each page"""
        canvas.saveState()
        
        # Footer with page number
        footer_text = f"Page {doc.page}"
        canvas.setFont('Helvetica', 9)
        canvas.setFillColor(colors.HexColor('#7f8c8d'))
        canvas.drawCentredString(
            letter[0] / 2,
            0.5 * inch,
            footer_text
        )
        
        # Header line
        canvas.setStrokeColor(colors.HexColor('#3498db'))
        canvas.setLineWidth(2)
        canvas.line(
            0.75 * inch,
            letter[1] - 0.75 * inch,
            letter[0] - 0.75 * inch,
            letter[1] - 0.75 * inch
        )
        
        canvas.restoreState()
    
    def generate(
        self,
        report_id: str,
        title: str,
        content: str,
        data: Dict[str, Any],
        output_dir: str
    ) -> str:
        """
        Generate PDF report
        
        Args:
            report_id: Unique report identifier
            title: Report title
            content: Rendered HTML content
            data: Report data dictionary
            output_dir: Output directory path
            
        Returns:
            Path to generated PDF file
        """
        # Create output file path
        filename = f"report_{report_id}.pdf"
        filepath = os.path.join(output_dir, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=0.75*inch
        )
        
        # Build story (content elements)
        story = []
        
        # Add title
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 0.2*inch))
        
        # Add metadata
        date_str = datetime.now().strftime("%B %d, %Y")
        metadata = f"Generated: {date_str}"
        if data.get('data', {}).get('authors'):
            metadata += f" | Authors: {data['data']['authors']}"
        story.append(Paragraph(metadata, self.styles['Metadata']))
        story.append(Spacer(1, 0.3*inch))
        
        # Parse and add sections
        sections = data.get('sections', [])
        for section in sections:
            # Add section heading
            story.append(Paragraph(section.title, self.styles['CustomHeading']))
            story.append(Spacer(1, 0.1*inch))
            
            # Add section content
            if section.content_type == 'text':
                content_text = section.content or data.get('data', {}).get(section.title.lower(), 'Content not provided.')
                # Clean HTML tags for PDF
                clean_text = self._clean_html(content_text)
                story.append(Paragraph(clean_text, self.styles['CustomBody']))
                story.append(Spacer(1, 0.2*inch))
            
            elif section.content_type == 'methods' and data.get('methods'):
                methods_text = self._clean_html(data['methods'])
                story.append(Paragraph(methods_text, self.styles['CustomBody']))
                story.append(Spacer(1, 0.2*inch))
            
            elif section.content_type == 'table':
                table_data = data.get('data', {}).get('tables', {}).get(section.title)
                if table_data:
                    # Create table
                    table_content = [table_data.get('headers', [])]
                    table_content.extend(table_data.get('rows', []))
                    
                    t = Table(table_content)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.2*inch))
            
            elif section.content_type == 'figure':
                figure_data = data.get('data', {}).get('figures', {}).get(section.title)
                if figure_data and os.path.exists(figure_data.get('path', '')):
                    # Add image
                    img = Image(figure_data['path'], width=5*inch, height=3*inch)
                    story.append(img)
                    # Add caption
                    caption = figure_data.get('caption', '')
                    story.append(Paragraph(f"<i>{caption}</i>", self.styles['CustomBody']))
                    story.append(Spacer(1, 0.2*inch))
        
        # Add references if available
        citations = data.get('citations', [])
        if citations:
            story.append(PageBreak())
            story.append(Paragraph("References", self.styles['CustomHeading']))
            story.append(Spacer(1, 0.1*inch))
            for i, citation in enumerate(citations, 1):
                story.append(Paragraph(f"{i}. {citation}", self.styles['CustomBody']))
        
        # Build PDF
        doc.build(story, onFirstPage=self._add_header_footer, onLaterPages=self._add_header_footer)
        
        return filepath
    
    def _clean_html(self, html_text: str) -> str:
        """Remove HTML tags from text"""
        # Simple HTML tag removal
        clean = re.sub('<[^<]+?>', '', html_text)
        return clean.strip()


class DOCXGenerator:
    """
    Word document generator using python-docx
    
    Requirements: 5.1
    """
    
    def __init__(self):
        """Initialize DOCX generator"""
        pass
    
    def generate(
        self,
        report_id: str,
        title: str,
        content: str,
        data: Dict[str, Any],
        output_dir: str
    ) -> str:
        """
        Generate DOCX report
        
        Args:
            report_id: Unique report identifier
            title: Report title
            content: Rendered HTML content
            data: Report data dictionary
            output_dir: Output directory path
            
        Returns:
            Path to generated DOCX file
        """
        # Create output file path
        filename = f"report_{report_id}.docx"
        filepath = os.path.join(output_dir, filename)
        
        # Create document
        doc = Document()
        
        # Set document styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        date_str = datetime.now().strftime("%B %d, %Y")
        metadata = f"Generated: {date_str}"
        if data.get('data', {}).get('authors'):
            metadata += f"\nAuthors: {data['data']['authors']}"
        meta_para = doc.add_paragraph(metadata)
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.runs[0]
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()  # Spacer
        
        # Add sections
        sections = data.get('sections', [])
        for section in sections:
            # Add section heading
            doc.add_heading(section.title, level=1)
            
            # Add section content
            if section.content_type == 'text':
                content_text = section.content or data.get('data', {}).get(section.title.lower(), 'Content not provided.')
                clean_text = self._clean_html(content_text)
                para = doc.add_paragraph(clean_text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            elif section.content_type == 'methods' and data.get('methods'):
                methods_text = self._clean_html(data['methods'])
                para = doc.add_paragraph(methods_text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            elif section.content_type == 'table':
                table_data = data.get('data', {}).get('tables', {}).get(section.title)
                if table_data:
                    headers = table_data.get('headers', [])
                    rows = table_data.get('rows', [])
                    
                    # Create table
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add headers
                    for i, header in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = str(header)
                        cell.paragraphs[0].runs[0].font.bold = True
                    
                    # Add rows
                    for row_idx, row in enumerate(rows, 1):
                        for col_idx, cell_value in enumerate(row):
                            table.rows[row_idx].cells[col_idx].text = str(cell_value)
            
            elif section.content_type == 'figure':
                figure_data = data.get('data', {}).get('figures', {}).get(section.title)
                if figure_data and os.path.exists(figure_data.get('path', '')):
                    doc.add_picture(figure_data['path'], width=Inches(5))
                    caption = figure_data.get('caption', '')
                    caption_para = doc.add_paragraph(caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.runs[0]
                    caption_run.italic = True
            
            doc.add_paragraph()  # Spacer
        
        # Add references
        citations = data.get('citations', [])
        if citations:
            doc.add_page_break()
            doc.add_heading('References', level=1)
            for i, citation in enumerate(citations, 1):
                doc.add_paragraph(f"{i}. {citation}")
        
        # Save document
        doc.save(filepath)
        
        return filepath
    
    def _clean_html(self, html_text: str) -> str:
        """Remove HTML tags from text"""
        clean = re.sub('<[^<]+?>', '', html_text)
        return clean.strip()


class LaTeXGenerator:
    """
    LaTeX document generator using PyLaTeX
    
    Requirements: 5.1
    """
    
    def __init__(self):
        """Initialize LaTeX generator"""
        pass
    
    def generate(
        self,
        report_id: str,
        title: str,
        content: str,
        data: Dict[str, Any],
        output_dir: str
    ) -> str:
        """
        Generate LaTeX report
        
        Args:
            report_id: Unique report identifier
            title: Report title
            content: Rendered HTML content
            data: Report data dictionary
            output_dir: Output directory path
            
        Returns:
            Path to generated LaTeX file
        """
        # Create output file path
        filename = f"report_{report_id}.tex"
        filepath = os.path.join(output_dir, filename)
        
        # Create LaTeX document
        doc = LaTeXDocument()
        
        # Add packages
        doc.packages.append(Package('geometry', options=['margin=1in']))
        doc.packages.append(Package('graphicx'))
        doc.packages.append(Package('booktabs'))
        doc.packages.append(Package('hyperref'))
        doc.packages.append(Package('natbib'))
        
        # Add title
        doc.preamble.append(Command('title', title))
        doc.preamble.append(Command('date', NoEscape(r'\today')))
        
        # Add author if available
        if data.get('data', {}).get('authors'):
            doc.preamble.append(Command('author', data['data']['authors']))
        
        doc.append(NoEscape(r'\maketitle'))
        doc.append(NoEscape(r'\tableofcontents'))
        doc.append(NoEscape(r'\newpage'))
        
        # Add sections
        sections = data.get('sections', [])
        for section in sections:
            with doc.create(Section(section.title)):
                if section.content_type == 'text':
                    content_text = section.content or data.get('data', {}).get(section.title.lower(), 'Content not provided.')
                    clean_text = self._clean_html(content_text)
                    doc.append(clean_text)
                
                elif section.content_type == 'methods' and data.get('methods'):
                    methods_text = self._clean_html(data['methods'])
                    doc.append(methods_text)
                
                elif section.content_type == 'table':
                    table_data = data.get('data', {}).get('tables', {}).get(section.title)
                    if table_data:
                        # LaTeX table generation would go here
                        doc.append("Table content placeholder")
                
                elif section.content_type == 'figure':
                    figure_data = data.get('data', {}).get('figures', {}).get(section.title)
                    if figure_data and os.path.exists(figure_data.get('path', '')):
                        # LaTeX figure inclusion would go here
                        doc.append("Figure placeholder")
        
        # Add bibliography
        citations = data.get('citations', [])
        if citations:
            with doc.create(Section('References')):
                doc.append(NoEscape(r'\begin{enumerate}'))
                for citation in citations:
                    doc.append(NoEscape(f'\\item {citation}'))
                doc.append(NoEscape(r'\end{enumerate}'))
        
        # Generate LaTeX file
        doc.generate_tex(filepath.replace('.tex', ''))
        
        return filepath
    
    def _clean_html(self, html_text: str) -> str:
        """Remove HTML tags and escape LaTeX special characters"""
        # Remove HTML tags
        clean = re.sub('<[^<]+?>', '', html_text)
        # Escape LaTeX special characters
        special_chars = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\^{}',
            '\\': r'\textbackslash{}'
        }
        for char, escaped in special_chars.items():
            clean = clean.replace(char, escaped)
        return clean.strip()
